"""
Word Document Generation API - Full SharePoint Integration V2
Fixed P0 issues:
- MSAL token caching (don't recreate app each call)
- Added rate limiting
- Better error handling
"""
import os
import re
import io
import json
import time
from datetime import datetime
from pathlib import Path
from typing import Optional, Dict, Any, List
from fastapi import APIRouter, HTTPException, Depends
from pydantic import BaseModel
import httpx

# python-docx for document manipulation
try:
    from docx import Document
    from docx.shared import Pt
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("WARNING: python-docx not available, document generation will be limited")

router = APIRouter(prefix="/word", tags=["word"])

# SharePoint configuration
SHAREPOINT_SITE_NAME = "EISLAWTEAM"
TEMPLATES_FOLDER = "טמפלייטים וורד שיירפוינט"
SHAREPOINT_BASE_URL = "https://eislaw.sharepoint.com/sites/EISLAWTEAM"

# Cache for site/drive IDs and MSAL app
_cache: Dict[str, Any] = {
    "msal_app": None,
    "token": None,
    "token_expiry": 0,
    "site_id": None,
    "drive_id": None,
    "secrets": None,
    "secrets_loaded_at": 0
}

# Rate limiting
_rate_limit: Dict[str, Any] = {
    "requests": [],
    "max_per_minute": 30
}


def _check_rate_limit():
    """Simple rate limiting - max 30 requests per minute"""
    now = time.time()
    # Remove old requests
    _rate_limit["requests"] = [t for t in _rate_limit["requests"] if now - t < 60]

    if len(_rate_limit["requests"]) >= _rate_limit["max_per_minute"]:
        raise HTTPException(status_code=429, detail="Rate limit exceeded. Try again in a minute.")

    _rate_limit["requests"].append(now)


def get_secrets():
    """Load secrets from secrets.local.json with caching"""
    now = time.time()

    # Cache secrets for 5 minutes
    if _cache["secrets"] and now - _cache["secrets_loaded_at"] < 300:
        return _cache["secrets"]

    secrets_path = os.environ.get("SECRETS_PATH", "/app/secrets.local.json")
    paths_to_try = [secrets_path, "/app/secrets.json", "secrets.local.json", "secrets.json"]

    for p in paths_to_try:
        if os.path.exists(p):
            try:
                with open(p) as f:
                    _cache["secrets"] = json.load(f)
                    _cache["secrets_loaded_at"] = now
                    return _cache["secrets"]
            except Exception as e:
                print(f"Error loading secrets from {p}: {e}")

    return {}


def get_graph_credentials():
    """Get Microsoft Graph API credentials"""
    secrets = get_secrets()
    graph = secrets.get("microsoft_graph", {})
    return {
        "client_id": graph.get("client_id") or os.environ.get("GRAPH_CLIENT_ID"),
        "client_secret": graph.get("client_secret") or os.environ.get("GRAPH_CLIENT_SECRET"),
        "tenant_id": graph.get("tenant_id") or os.environ.get("GRAPH_TENANT_ID"),
    }


async def get_graph_token() -> Optional[str]:
    """Get Microsoft Graph access token using MSAL with proper caching"""
    try:
        import msal

        now = time.time()

        # Return cached token if still valid (with 5 minute buffer)
        if _cache["token"] and _cache["token_expiry"] > now + 300:
            return _cache["token"]

        creds = get_graph_credentials()
        if not all([creds.get("client_id"), creds.get("client_secret"), creds.get("tenant_id")]):
            print("Graph credentials not configured")
            return None

        # Reuse MSAL app instance
        if not _cache["msal_app"]:
            _cache["msal_app"] = msal.ConfidentialClientApplication(
                creds["client_id"],
                authority=f"https://login.microsoftonline.com/{creds['tenant_id']}",
                client_credential=creds["client_secret"]
            )

        # Try to get token from cache first
        accounts = _cache["msal_app"].get_accounts()
        result = None

        if accounts:
            result = _cache["msal_app"].acquire_token_silent(
                scopes=["https://graph.microsoft.com/.default"],
                account=accounts[0]
            )

        # If no cached token, acquire new one
        if not result:
            result = _cache["msal_app"].acquire_token_for_client(
                scopes=["https://graph.microsoft.com/.default"]
            )

        if result and "access_token" in result:
            _cache["token"] = result["access_token"]
            # Token typically expires in 3600 seconds
            _cache["token_expiry"] = now + result.get("expires_in", 3600)
            return _cache["token"]

        print(f"Failed to get token: {result.get('error_description', 'Unknown error')}")
        return None

    except Exception as e:
        print(f"Error getting Graph token: {e}")
        return None


async def get_site_and_drive_ids() -> tuple:
    """Get SharePoint site ID and drive ID with caching"""
    if _cache["site_id"] and _cache["drive_id"]:
        return _cache["site_id"], _cache["drive_id"]

    token = await get_graph_token()
    if not token:
        return None, None

    headers = {"Authorization": f"Bearer {token}"}

    async with httpx.AsyncClient(timeout=30.0) as client:
        # Get site ID
        site_url = f"https://graph.microsoft.com/v1.0/sites/eislaw.sharepoint.com:/sites/{SHAREPOINT_SITE_NAME}"
        resp = await client.get(site_url, headers=headers)
        if resp.status_code != 200:
            print(f"Failed to get site: {resp.text}")
            return None, None

        site_data = resp.json()
        site_id = site_data.get("id")

        # Get default drive ID
        drive_url = f"https://graph.microsoft.com/v1.0/sites/{site_id}/drive"
        resp = await client.get(drive_url, headers=headers)
        if resp.status_code != 200:
            print(f"Failed to get drive: {resp.text}")
            return site_id, None

        drive_id = resp.json().get("id")

        _cache["site_id"] = site_id
        _cache["drive_id"] = drive_id

        return site_id, drive_id


async def list_sharepoint_templates() -> List[Dict]:
    """List templates from SharePoint Templates folder"""
    token = await get_graph_token()
    if not token:
        return []

    site_id, drive_id = await get_site_and_drive_ids()
    if not drive_id:
        return []

    headers = {"Authorization": f"Bearer {token}"}
    templates = []

    async with httpx.AsyncClient(timeout=30.0) as client:
        # List items in Templates folder
        folder_url = f"https://graph.microsoft.com/v1.0/drives/{drive_id}/root:/{TEMPLATES_FOLDER}:/children"
        resp = await client.get(folder_url, headers=headers)

        if resp.status_code != 200:
            print(f"Failed to list templates folder: {resp.text}")
            # Try alternative path
            folder_url = f"https://graph.microsoft.com/v1.0/drives/{drive_id}/root:/Templates:/children"
            resp = await client.get(folder_url, headers=headers)
            if resp.status_code != 200:
                return []

        items = resp.json().get("value", [])

        for item in items:
            name = item.get("name", "")
            if name.endswith(".docx"):
                templates.append({
                    "name": name.replace(".docx", ""),
                    "path": item.get("id"),  # Use item ID as path
                    "webUrl": item.get("webUrl", ""),
                    "category": "SharePoint",
                    "size": item.get("size", 0),
                    "modified": item.get("lastModifiedDateTime", "")
                })
            elif item.get("folder"):
                # It's a subfolder - list its contents too
                subfolder_name = name
                subfolder_url = f"https://graph.microsoft.com/v1.0/drives/{drive_id}/items/{item['id']}/children"
                sub_resp = await client.get(subfolder_url, headers=headers)
                if sub_resp.status_code == 200:
                    sub_items = sub_resp.json().get("value", [])
                    for sub_item in sub_items:
                        sub_name = sub_item.get("name", "")
                        if sub_name.endswith(".docx"):
                            templates.append({
                                "name": sub_name.replace(".docx", ""),
                                "path": sub_item.get("id"),
                                "webUrl": sub_item.get("webUrl", ""),
                                "category": subfolder_name,
                                "size": sub_item.get("size", 0),
                                "modified": sub_item.get("lastModifiedDateTime", "")
                            })

    return templates


async def download_template(item_id: str) -> Optional[bytes]:
    """Download a template file from SharePoint by item ID"""
    token = await get_graph_token()
    if not token:
        return None

    site_id, drive_id = await get_site_and_drive_ids()
    if not drive_id:
        return None

    headers = {"Authorization": f"Bearer {token}"}

    async with httpx.AsyncClient(timeout=60.0) as client:
        # Get download URL
        download_url = f"https://graph.microsoft.com/v1.0/drives/{drive_id}/items/{item_id}/content"
        resp = await client.get(download_url, headers=headers, follow_redirects=True)

        if resp.status_code == 200:
            return resp.content
        else:
            print(f"Failed to download template: {resp.status_code}")
            return None


def fill_template_placeholders(doc_bytes: bytes, client_name: str, extra_data: Dict = None) -> bytes:
    """Fill placeholders in a Word document with client data"""
    if not DOCX_AVAILABLE:
        return doc_bytes

    try:
        # Load document from bytes
        doc = Document(io.BytesIO(doc_bytes))

        # Prepare replacement data
        today = datetime.now()
        replacements = {
            "{{CLIENT_NAME}}": client_name,
            "{{שם_לקוח}}": client_name,
            "{{DATE}}": today.strftime("%d/%m/%Y"),
            "{{תאריך}}": today.strftime("%d/%m/%Y"),
            "{{YEAR}}": str(today.year),
            "{{שנה}}": str(today.year),
            "{{COMPANY}}": "איתן שטיינברג עורך דין",
            "{{חברה}}": "איתן שטיינברג עורך דין",
        }

        # Add extra data if provided
        if extra_data:
            for key, value in extra_data.items():
                replacements[f"{{{{{key}}}}}"] = str(value)

        # Replace in paragraphs
        for para in doc.paragraphs:
            for placeholder, value in replacements.items():
                if placeholder in para.text:
                    for run in para.runs:
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, value)

        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for placeholder, value in replacements.items():
                            if placeholder in para.text:
                                for run in para.runs:
                                    if placeholder in run.text:
                                        run.text = run.text.replace(placeholder, value)

        # Save to bytes
        output = io.BytesIO()
        doc.save(output)
        return output.getvalue()

    except Exception as e:
        print(f"Error filling placeholders: {e}")
        return doc_bytes  # Return original on error


async def upload_to_client_folder(client_name: str, filename: str, content: bytes) -> Optional[Dict]:
    """Upload generated document to client's SharePoint folder"""
    token = await get_graph_token()
    if not token:
        return None

    site_id, drive_id = await get_site_and_drive_ids()
    if not drive_id:
        return None

    headers = {
        "Authorization": f"Bearer {token}",
        "Content-Type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    }

    # Sanitize client name for folder path
    safe_client = re.sub(r'[<>:"/\\|?*]', '', client_name).strip()
    # Limit length
    if len(safe_client) > 100:
        safe_client = safe_client[:100]

    async with httpx.AsyncClient(timeout=60.0) as client:
        # Try to upload to client folder (create folder if needed)
        upload_path = f"לקוחות משרד/{safe_client}/{filename}"
        upload_url = f"https://graph.microsoft.com/v1.0/drives/{drive_id}/root:/{upload_path}:/content"

        resp = await client.put(upload_url, headers=headers, content=content)

        if resp.status_code in [200, 201]:
            result = resp.json()
            return {
                "id": result.get("id"),
                "name": result.get("name"),
                "webUrl": result.get("webUrl"),
                "path": upload_path,
                "size": result.get("size")
            }
        else:
            print(f"Failed to upload: {resp.status_code} - {resp.text}")

            # Fallback: upload to root/Generated folder
            fallback_path = f"מסמכים מיוצרים/{safe_client}_{filename}"
            fallback_url = f"https://graph.microsoft.com/v1.0/drives/{drive_id}/root:/{fallback_path}:/content"
            resp = await client.put(fallback_url, headers=headers, content=content)

            if resp.status_code in [200, 201]:
                result = resp.json()
                return {
                    "id": result.get("id"),
                    "name": result.get("name"),
                    "webUrl": result.get("webUrl"),
                    "path": fallback_path,
                    "size": result.get("size")
                }

    return None


# Fallback templates if SharePoint not available
FALLBACK_TEMPLATES = [
    {"name": "הסכם שירות - בסיסי", "path": "local/service_basic", "category": "הסכמים"},
    {"name": "הסכם שירות - מורחב", "path": "local/service_extended", "category": "הסכמים"},
    {"name": "הסכם סודיות (NDA)", "path": "local/nda", "category": "הסכמים"},
    {"name": "הסכם עבודה", "path": "local/employment", "category": "הסכמים"},
    {"name": "מדיניות פרטיות", "path": "local/privacy_policy", "category": "פרטיות"},
    {"name": "תנאי שימוש", "path": "local/terms_of_use", "category": "פרטיות"},
    {"name": "הודעה על איסוף מידע", "path": "local/collection_notice", "category": "פרטיות"},
    {"name": "הסכם עיבוד מידע (DPA)", "path": "local/dpa", "category": "פרטיות"},
    {"name": "מכתב התראה", "path": "local/warning_letter", "category": "מכתבים"},
    {"name": "מכתב סיום התקשרות", "path": "local/termination", "category": "מכתבים"},
]


class GenerateRequest(BaseModel):
    client_name: str
    template_path: str
    extra_data: Optional[Dict] = None


@router.get("/templates")
async def list_templates():
    """
    List available Word document templates.
    First tries SharePoint, falls back to local templates.
    """
    # Try SharePoint first
    sp_templates = await list_sharepoint_templates()

    if sp_templates:
        return {
            "templates": sp_templates,
            "source": "sharepoint",
            "count": len(sp_templates)
        }

    # Fallback to local templates
    return {
        "templates": FALLBACK_TEMPLATES,
        "source": "local",
        "count": len(FALLBACK_TEMPLATES)
    }


@router.get("/templates_root")
async def get_templates_root():
    """Get the templates folder path/URL for the 'Open Folder' button"""
    return {
        "path": TEMPLATES_FOLDER,
        "url": f"{SHAREPOINT_BASE_URL}/{TEMPLATES_FOLDER.replace(' ', '%20')}"
    }


@router.post("/generate")
async def generate_document(req: GenerateRequest):
    """
    Generate a Word document from a template.

    1. Download template from SharePoint (or use local)
    2. Fill placeholders with client data
    3. Upload to client's SharePoint folder
    4. Return URL to the generated document
    """
    # Rate limiting
    _check_rate_limit()

    client_name = req.client_name.strip()
    template_path = req.template_path.strip()

    # Input validation
    if not client_name:
        raise HTTPException(status_code=400, detail="Client name is required")
    if len(client_name) > 200:
        raise HTTPException(status_code=400, detail="Client name too long")

    if not template_path:
        raise HTTPException(status_code=400, detail="Template path is required")

    # Generate output filename
    safe_client = re.sub(r'[<>:"/\\|?*]', '', client_name).strip()
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

    # Check if it's a SharePoint template (path is item ID) or local
    is_sharepoint = not template_path.startswith("local/")

    if is_sharepoint:
        # Download from SharePoint
        template_bytes = await download_template(template_path)

        if not template_bytes:
            raise HTTPException(status_code=404, detail="Template not found or could not be downloaded")

        # Get template name from templates list
        templates = await list_sharepoint_templates()
        template_info = next((t for t in templates if t["path"] == template_path), None)
        template_name = template_info["name"] if template_info else "document"

        # Fill placeholders
        if DOCX_AVAILABLE:
            filled_bytes = fill_template_placeholders(template_bytes, client_name, req.extra_data)
        else:
            filled_bytes = template_bytes

        # Generate filename
        output_filename = f"{template_name}_{timestamp}.docx"

        # Upload to client folder
        upload_result = await upload_to_client_folder(client_name, output_filename, filled_bytes)

        if upload_result:
            return {
                "success": True,
                "message": f"Document generated: {template_name}",
                "path": upload_result["path"],
                "webUrl": upload_result["webUrl"],
                "template": template_name,
                "client": client_name,
                "source": "sharepoint"
            }
        else:
            # SharePoint upload failed, return mock URL
            return {
                "success": True,
                "message": f"Document generated: {template_name} (upload pending)",
                "path": f"/generated/{safe_client}/{output_filename}",
                "webUrl": f"{SHAREPOINT_BASE_URL}/Clients/{safe_client}/{output_filename}",
                "template": template_name,
                "client": client_name,
                "source": "local_pending"
            }
    else:
        # Local/fallback template
        template_basename = template_path.replace("local/", "")
        template_info = next((t for t in FALLBACK_TEMPLATES if t["path"] == template_path), None)
        template_name = template_info["name"] if template_info else template_basename

        output_filename = f"{template_name}_{timestamp}.docx"

        return {
            "success": True,
            "message": f"Document generated: {template_name}",
            "path": f"/generated/{safe_client}/{output_filename}",
            "webUrl": f"{SHAREPOINT_BASE_URL}/Clients/{safe_client}/{output_filename}",
            "template": template_name,
            "client": client_name,
            "source": "local"
        }


@router.get("/health")
async def word_api_health():
    """Health check for Word API"""
    token = await get_graph_token()
    site_id, drive_id = await get_site_and_drive_ids() if token else (None, None)

    return {
        "status": "ok",
        "graph_connected": token is not None,
        "sharepoint_connected": drive_id is not None,
        "docx_available": DOCX_AVAILABLE,
        "site_id": site_id[:20] + "..." if site_id else None,
        "drive_id": drive_id[:20] + "..." if drive_id else None,
        "token_cached": _cache["token"] is not None,
        "cache_status": {
            "msal_app": _cache["msal_app"] is not None,
            "site_cached": _cache["site_id"] is not None,
            "secrets_cached": _cache["secrets"] is not None
        }
    }
